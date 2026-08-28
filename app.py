from flask import Flask, render_template, request, send_file, jsonify, flash
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from werkzeug.utils import secure_filename
from docx import Document
from docxcompose.composer import Composer
from PIL import Image

import os
import uuid
import re

app = Flask(__name__)
app.secret_key = "smartmerge_glassy_secret_key"

UPLOAD_FOLDER = "uploads"
MERGED_FOLDER = "merged"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(MERGED_FOLDER, exist_ok=True)

MAX_FILE_SIZE = 50 * 1024 * 1024
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def image_to_pdf(image_path, output_pdf_path):
    """Converts image file to PDF using Pillow"""
    img = Image.open(image_path)
    if img.mode != 'RGB':
        img = img.convert('RGB')
    img.save(output_pdf_path, 'PDF', resolution=100.0)
    return output_pdf_path


def txt_to_pdf(txt_path, output_pdf_path):
    """Converts a text file to a basic PDF using reportlab or PyPDF fallback"""
    # Simple line-by-line text to PDF or fallback handling
    pass


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/merger")
def merger_workspace():
    return render_template("merger.html")


@app.route("/split", methods=["GET", "POST"])
def split_pdf():
    if request.method == "POST":
        file = request.files.get("file")
        page_range = request.form.get("page_range", "").strip()

        if not file or file.filename == "":
            return jsonify({"error": "No file uploaded"}), 400

        if not file.filename.lower().endswith(".pdf"):
            return jsonify({"error": "Only PDF files can be split"}), 400

        filename = secure_filename(file.filename)
        save_path = os.path.join(UPLOAD_FOLDER, f"split_input_{uuid.uuid4()}_{filename}")
        file.save(save_path)

        try:
            reader = PdfReader(save_path)
            total_pages = len(reader.pages)
            writer = PdfWriter()

            # Parse page numbers/ranges e.g. "1-3, 5" (1-indexed input)
            selected_pages = []
            if page_range:
                parts = page_range.split(",")
                for part in parts:
                    part = part.strip()
                    if "-" in part:
                        start, end = part.split("-")
                        for p in range(int(start), int(end) + 1):
                            if 1 <= p <= total_pages:
                                selected_pages.append(p - 1)
                    elif part.isdigit():
                        p = int(part)
                        if 1 <= p <= total_pages:
                            selected_pages.append(p - 1)
            else:
                # Default: split all pages or keep first page
                selected_pages = list(range(total_pages))

            for p_idx in selected_pages:
                writer.add_page(reader.pages[p_idx])

            output_id = str(uuid.uuid4())
            out_filename = f"split_{output_id}.pdf"
            output_path = os.path.join(MERGED_FOLDER, out_filename)

            with open(output_path, "wb") as f_out:
                writer.write(f_out)

            return send_file(output_path, as_attachment=True, download_name=f"split_{filename}")

        except Exception as e:
            return jsonify({"error": f"Failed to split PDF: {str(e)}"}), 500
        finally:
            if os.path.exists(save_path):
                os.remove(save_path)

    return render_template("split.html")


@app.route("/features")
def features():
    return render_template("features.html")


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/merge", methods=["POST"])
def merge_files():
    uploaded_files = request.files.getlist("files")
    custom_name = request.form.get("custom_name", "").strip()

    if not uploaded_files or (len(uploaded_files) == 1 and uploaded_files[0].filename == ''):
        return jsonify({"error": "No files uploaded"}), 400

    pdf_merger = PdfMerger()
    docx_files = []
    temp_converted_pdfs = []

    output_id = str(uuid.uuid4())
    filename_base = secure_filename(custom_name) if custom_name else f"SmartMerge_{output_id[:8]}"
    if not filename_base:
        filename_base = f"SmartMerge_{output_id[:8]}"

    merged_pdf_path = os.path.join(MERGED_FOLDER, f"{filename_base}.pdf")
    merged_docx_path = os.path.join(MERGED_FOLDER, f"{filename_base}.docx")

    try:
        for file in uploaded_files:
            if file.filename == '':
                continue

            if not allowed_file(file.filename):
                continue

            orig_filename = secure_filename(file.filename)
            unique_filename = f"{uuid.uuid4()}_{orig_filename}"
            save_path = os.path.join(UPLOAD_FOLDER, unique_filename)
            file.save(save_path)

            extension = orig_filename.rsplit('.', 1)[1].lower()

            if extension == "pdf":
                pdf_merger.append(save_path)
            elif extension in ["png", "jpg", "jpeg"]:
                converted_pdf = save_path + ".pdf"
                image_to_pdf(save_path, converted_pdf)
                temp_converted_pdfs.append(converted_pdf)
                pdf_merger.append(converted_pdf)
            elif extension == "docx":
                docx_files.append(save_path)
            elif extension == "txt":
                # Convert txt to docx or handle text
                doc = Document()
                with open(save_path, "r", encoding="utf-8", errors="ignore") as f:
                    doc.add_paragraph(f.read())
                txt_docx_path = save_path + ".docx"
                doc.save(txt_docx_path)
                docx_files.append(txt_docx_path)

        # Merge PDFs if pages added
        pdf_success = False
        if len(pdf_merger.pages) > 0:
            pdf_merger.write(merged_pdf_path)
            pdf_merger.close()
            pdf_success = True

        # Merge DOCX files
        docx_success = False
        if len(docx_files) > 0:
            master = Document(docx_files[0])
            composer = Composer(master)
            for doc in docx_files[1:]:
                composer.append(Document(doc))
            composer.save(merged_docx_path)
            docx_success = True

        # Return merged PDF first if present, otherwise DOCX
        if pdf_success and os.path.exists(merged_pdf_path):
            return send_file(merged_pdf_path, as_attachment=True, download_name=f"{filename_base}.pdf")
        elif docx_success and os.path.exists(merged_docx_path):
            return send_file(merged_docx_path, as_attachment=True, download_name=f"{filename_base}.docx")

        return jsonify({"error": "No valid files could be merged"}), 400

    except Exception as e:
        return jsonify({"error": f"Merge error: {str(e)}"}), 500

    finally:
        # Cleanup temporary files
        for temp_pdf in temp_converted_pdfs:
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)


if __name__ == "__main__":
    app.run(debug=True, port=5000)